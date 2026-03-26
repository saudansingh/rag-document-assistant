"""
Word document processor using python-docx
Extracts text and metadata from DOCX files
"""

from docx import Document as DocxDocument
from datetime import datetime
from typing import List, Dict, Any
import os


class DocxProcessor:
    """Handles DOCX text extraction with metadata using python-docx"""
    
    def __init__(self):
        self.supported_extensions = ['docx', 'doc']
    
    def extract_text_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from DOCX file with rich metadata
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of dictionaries containing text content and metadata
        """
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Document-level metadata
            core_props = doc.core_properties
            doc_metadata = {
                "source": file_path,
                "file_type": "docx",
                "processed_at": datetime.now().isoformat(),
                "title": core_props.title or '',
                "author": core_props.author or '',
                "subject": core_props.subject or '',
                "created": str(core_props.created) if core_props.created else '',
                "modified": str(core_props.modified) if core_props.modified else '',
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables)
            }
            
            # Extract text from paragraphs
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    # Paragraph-specific metadata
                    para_metadata = doc_metadata.copy()
                    para_metadata.update({
                        "paragraph_number": para_num + 1,
                        "paragraph_size": len(paragraph.text),
                        "style": paragraph.style.name if paragraph.style else 'Normal',
                        "alignment": str(paragraph.alignment) if paragraph.alignment else 'Unknown'
                    })
                    
                    text_content.append({
                        "content": paragraph.text,
                        "metadata": para_metadata
                    })
            
            # Extract text from tables
            for table_num, table in enumerate(doc.tables):
                for row_num, row in enumerate(table.rows):
                    for cell_num, cell in enumerate(row.cells):
                        if cell.text.strip():
                            table_metadata = doc_metadata.copy()
                            table_metadata.update({
                                "content_type": "table",
                                "table_number": table_num + 1,
                                "row_number": row_num + 1,
                                "cell_number": cell_num + 1,
                                "cell_size": len(cell.text)
                            })
                            
                            text_content.append({
                                "content": cell.text,
                                "metadata": table_metadata
                            })
            
            return text_content
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing DOCX metadata
        """
        try:
            doc = DocxDocument(file_path)
            core_props = doc.core_properties
            
            return {
                "title": core_props.title or '',
                "author": core_props.author or '',
                "subject": core_props.subject or '',
                "created": str(core_props.created) if core_props.created else '',
                "modified": str(core_props.modified) if core_props.modified else '',
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
                "sections_count": len(doc.sections)
            }
            
        except Exception as e:
            raise Exception(f"Error extracting DOCX metadata: {str(e)}")
    
    def is_supported(self, file_extension: str) -> bool:
        """Check if file extension is supported"""
        return file_extension.lower() in self.supported_extensions
